"""
	This is the Report Creator for the Back SEO Application. This will module
	creates the base report for SEO agencies to use to send to their clients.
	The base report includes:
		- Company Headers
		- Company Logo
		- Company Contact Information

"""
from docx import Document
from docx.shared import Inches, Pt
from docx.document import Document as _Document, Section
from docx.text.paragraph import Paragraph, Run
from docx.shape import InlineShape
from docxtpl import DocxTemplate, InlineImage, Subdoc
from BackSEOSettings import BackSEOSettings, user_local_dir
import random
import os
import time
import pickle

from typing import List, Dict


def cleanFilename(s: str):
    if not s:
        return ""
    badchars = "\\/:*?\"'<>|&"
    for c in badchars:
        s = s.replace(c, "")
    return s


class SEOEmployee:
    def __init__(
        self,
        employee_name: str = "Employee Name",
        employee_position: str = "Employee Position",
        employee_phone: str = "123-456-7890",
        employee_email: str = "employee@companysite.com",
    ) -> None:
        self.employee_name: str = employee_name
        self.employee_position: str = employee_position
        self.employee_phone: str = employee_phone
        self.employee_email: str = employee_email
        self.outputFolder: str = f"SEO/Resources/employees"
        self.outputPath: str = f"SEO/Resources/employees/{self.employee_name}"

    def save(self) -> None:
        if not os.path.exists(self.outputFolder):
            os.makedirs(self.outputFolder)
        with open(self.outputPath, "wb") as f:
            pickle.dump(self, f, pickle.HIGHEST_PROTOCOL)


class SEOClient:
    """- Client Name
    - Client Address
    - Client Phone Number
    - Client Logo
    - Client Website
    - Client Social Media
    - Client Target Keywords
    - Client Target Locations
    """

    def __init__(
        self,
        client_name: str = "Client Name",
        client_address: str = "1234 Client St. City, State 12345",
        client_phone: str = "123-456-7890",
        client_logo: str = "clientLogo.png",
        client_email: str = "client@clientsite.com",
        client_website: str = "clientsite.com",
        client_website_sitemap: str = "https://clientsite.com/sitemap.xml",
        client_social_media: Dict[str, str] = {},
        client_target_keywords: List[str] = [],
        client_target_locations: List[str] = [],
    ) -> None:
        self.client_name: str = client_name
        self.client_address: str = client_address
        self.client_phone: str = client_phone
        self.client_logo: str = client_logo
        self.client_email: str = client_email
        self.client_website: str = client_website
        self.client_website_sitemap: str = client_website_sitemap
        self.client_social_media: Dict[str, str] = client_social_media
        self.client_target_keywords: List[str] = client_target_keywords
        self.client_target_locations: List[str] = client_target_locations
        self.rankingsOverTime: dict[
            str, dict[str, dict[str, int]]
        ] = dict()  # {keyword: {location: {date: rank}}}

    def addRankings(self, keyword: str, location: str, date: str, rank: int) -> None:
        if keyword not in self.rankingsOverTime:
            self.rankingsOverTime[keyword] = dict()
        if location not in self.rankingsOverTime[keyword]:
            self.rankingsOverTime[keyword][location] = dict()
        self.rankingsOverTime[keyword][location][date] = rank


class SEOAgency:
    def __init__(
        self,
        company_name: str = "Back SEO Marketing Software",
        company_address: str = "1234 Main St. City, State 12345",
        company_phone: str = "123-456-7890",
        company_email: str = "outreach@backseo.org",
        company_website: str = "backseo.org",
        company_logo: str = "bSEOLogo.png",
    ):
        """Creates a new SEO Agency object. This is used to create:
        - SEO Competitor Analysis Reports
        - SEO Website Audit Reports
        - Content Analysis Reports
        - SEO Keyword Research Reports
        - Progress Updates
        """
        self.company_name: str = company_name
        self.company_address: str = company_address
        self.company_phone: str = company_phone
        self.company_email: str = company_email
        self.company_website: str = company_website
        self.company_logo: str = company_logo
        self.output_directory: str = user_local_dir().as_posix()
        self.employees: List[SEOEmployee] = list()
        self.clients: Dict[str, SEOClient] = dict()
        self.socialMedia: Dict[str, str] = dict()

    def save(self) -> None:
        if not os.path.exists(user_local_dir()):
            os.makedirs(user_local_dir())
        with open(user_local_dir() / "seo.agency", "wb") as f:
            pickle.dump(self, f, pickle.HIGHEST_PROTOCOL)

    def addClient(self, client: SEOClient) -> None:
        self.clients[client.client_name] = client

    def createClientRankingSheets(self) -> None:
        for name, client in self.clients.items():
            try:
                with open(
                    f"{self.output_directory}/reports/{name}/rankings.csv", "w"
                ) as f:
                    f.write("Keyword,Location,Date,Rank\n")
                    for keyword, locations in client.rankingsOverTime.items():
                        for location, dates in locations.items():
                            for date, rank in dates.items():
                                f.write(f"{keyword},{location},{date},{rank}\n")
            except:
                continue


class SEOReport:
    def __init__(
        self,
        agency: SEOAgency,
        title: str = "Local SEO Analysis Report",
        subtitle: str = "Site Audit & Competition",
        coverImage: str = "agency reports1000x2.png",
        footerLeft: str = "Report for CLIENT",
        footerCenter: str = "Your Key to Success",
        footerRight: str = "View Report -->",
    ) -> None:
        """Creates a new SEO Report object. This is used to create:
        - SEO Competitor Analysis Reports
        - SEO Website Audit Reports
        - Content Analysis Reports
        - SEO Keyword Research Reports
        - Progress Updates
        """
        self.agency: SEOAgency = agency
        self.title: str = title
        self.subtitle: str = subtitle
        if os.path.exists(coverImage):
            self.coverImage: str = coverImage
        else:
            self.coverImage = "agency reports.png"
        self.document: _Document = Document()  # type; _Document
        self.section: list[Section] | Section = self.document.sections[0]
        self.footerLeft: str = footerLeft
        self.footerCenter: str = footerCenter
        self.footerRight: str = footerRight

    def coverPage(self) -> None:
        self.doctitle: Paragraph = self.document.add_heading(self.title, 0)

        self.bigcoverpic: Paragraph = self.document.add_paragraph()
        self.bigcoverpic1: Run = self.bigcoverpic.add_run()
        self.bigcoverpic2: InlineShape = self.bigcoverpic1.add_picture(
            self.coverImage, width=Inches(4), height=Inches(4)
        )

        self.docsubtitle: Paragraph = self.document.add_heading(self.subtitle, 1)
        self.coverPic: Paragraph = self.document.add_paragraph()
        self.coverPic.add_run().add_picture(
            self.agency.company_logo, width=Inches(1.25)
        )
        self.coverPic.alignment = 1
        company_name: Paragraph = self.document.add_paragraph(self.agency.company_name)
        company_address: Paragraph = self.document.add_paragraph(
            self.agency.company_address
        )
        company_phone: Paragraph = self.document.add_paragraph(
            self.agency.company_phone
        )
        company_email: Paragraph = self.document.add_paragraph(
            self.agency.company_email
        )
        company_website: Paragraph = self.document.add_paragraph(
            self.agency.company_website
        )
        company_name.alignment = 1
        company_address.alignment = 1
        company_phone.alignment = 1
        company_email.alignment = 1
        company_website.alignment = 1
        company_name.paragraph_format.keep_with_next = True
        company_address.paragraph_format.keep_with_next = True
        company_phone.paragraph_format.keep_with_next = True
        company_email.paragraph_format.keep_with_next = True
        company_website.paragraph_format.keep_with_next = True
        self.doctitle.alignment = 1
        self.bigcoverpic.alignment = 1
        self.docsubtitle.alignment = 1
        self.document.add_page_break()

    def coverHeaderFooter(self) -> None:
        self.section.different_first_page_header_footer = True
        footer = self.section.first_page_footer
        header = self.section.first_page_header
        hp = header.paragraphs[0]
        hp.clear().add_run().add_picture(self.agency.company_logo, width=Inches(0.25))
        hp.add_run(f"\t\t{self.agency.company_phone}")
        fp = footer.paragraphs[0]
        fp.clear().add_run(
            f"{self.footerLeft}\t{self.footerCenter}\t{self.footerRight}"
        )

    def save(self, filename: str = "SEO Report") -> None:
        self.fileName: str = f"output/reports/{cleanFilename(filename)}{random.randint(0,1000)}.docx"
        self.document.save(self.fileName)


def createLocalReport(auddict: dict) -> str | None:
    context = auddict.copy()
    basiccontext = dict()
    theBusiness = auddict["clientName"]
    keyword = auddict["keyword"]
    basiccontext["clientName"] = auddict["clientName"]
    basiccontext["keyword"] = auddict["keyword"]
    basiccontext["zoom"] = auddict["zoom"]

    aOut: str = BackSEOSettings().agencyOutput
    outpath: str = (
        f"{aOut}/reports/{cleanFilename(theBusiness)}/local/{cleanFilename(keyword)}"
    )
    if not os.path.exists(outpath):
        os.makedirs(outpath)
    tpl = DocxTemplate("SEO/Resources/templates/local_audit_template.docx")
    basictpl = DocxTemplate("SEO/Resources/templates/local_audit_template_basic.docx")
    basiccontext["mapImage"] = InlineImage(
        basictpl, auddict["mapImage"], width=Inches(6.5)
    )
    basiccontext["clientImage"] = InlineImage(
        basictpl, auddict["clientImage"], height=Inches(2.5), width=Inches(6.5)
    )
    context["mapImage"] = InlineImage(tpl, auddict["mapImage"], width=Inches(6.5))
    h = 0.8
    w = 1.73
    context["clientImage"] = InlineImage(
        tpl, auddict["clientImage"], height=Inches(2.5), width=Inches(6.5)
    )
    directions = auddict["tdirections"]
    directions2 = list()
    basicdirections = list()
    for direction in directions:
        basicdirection = dict()
        if os.path.exists(direction["rank1"]):
            basicdirection["rank1"] = InlineImage(
                basictpl, direction["rank1"], height=Inches(h), width=Inches(w)
            )
            direction["rank1"] = InlineImage(
                tpl, direction["rank1"], height=Inches(h), width=Inches(w)
            )
        else:
            basicdirection["rank1"] = InlineImage(
                basictpl, "blank.png", height=Inches(h), width=Inches(w)
            )
            direction["rank1"] = InlineImage(
                tpl, "blank.png", height=Inches(h), width=Inches(w)
            )
        if os.path.exists(direction["rank2"]):
            basicdirection["rank2"] = InlineImage(
                basictpl, direction["rank2"], height=Inches(h), width=Inches(w)
            )
            direction["rank2"] = InlineImage(
                tpl, direction["rank2"], height=Inches(h), width=Inches(w)
            )
        else:
            basicdirection["rank2"] = InlineImage(
                basictpl, "blank.png", height=Inches(h), width=Inches(w)
            )
            direction["rank2"] = InlineImage(
                tpl, "blank.png", height=Inches(h), width=Inches(w)
            )
        if os.path.exists(direction["rank3"]):
            basicdirection["rank3"] = InlineImage(
                basictpl, direction["rank3"], height=Inches(h), width=Inches(w)
            )
            direction["rank3"] = InlineImage(
                tpl, direction["rank3"], height=Inches(h), width=Inches(w)
            )
        else:
            basicdirection["rank3"] = InlineImage(
                basictpl, "blank.png", height=Inches(h), width=Inches(w)
            )
            direction["rank3"] = InlineImage(
                tpl, "blank.png", height=Inches(h), width=Inches(w)
            )
        i = 1
        basicdirection["direction"] = direction["direction"]
        basicdirections.append(basicdirection)
        if "all" in direction:
            direction2 = dict()
            direction2["direction"] = direction["direction"]
            direction2["businesses"] = list()
            bus: str
            for bus in direction["all"]:
                name = bus.split("_")[-1].split(".")[0]
                if os.path.exists(bus):
                    direction2["businesses"].append(
                        {
                            "rank": f"{i}",
                            "name": name,
                            "image": InlineImage(
                                tpl, bus, height=Inches(h / 2), width=Inches(w / 2)
                            ),
                        }
                    )
                else:
                    direction2["businesses"].append(
                        {
                            "rank": f"{i}",
                            "name": name,
                            "image": InlineImage(
                                tpl,
                                "blank.png",
                                height=Inches(h / 2),
                                width=Inches(w / 2),
                            ),
                        }
                    )
                if theBusiness in name:
                    break
                i += 1
            directions2.append(direction2)
    context["tdirections2"] = directions2
    basiccontext["tdirections"] = basicdirections
    context["tdirections"] = directions
    context["zoom"] = auddict["zoom"]
    auditName: str = str(
        f"{outpath}/{str(auddict['zoom'])}_{time.strftime('%b_%d_%y_%H%M%S', time.gmtime())}.docx"
    )
    basicName: str = str(
        f"{outpath}/{str(auddict['zoom'])}_{time.strftime('%b_%d_%y_%H%M%S', time.gmtime())}basic.docx"
    )
    try:
        tpl.render(context, autoescape=True)
        tpl.save(auditName)
        basictpl.render(basiccontext, autoescape=True)
        basictpl.save(basicName)
        return auditName
    except Exception as e:
        print(e)
        return None


def createMasterReport(
    agency: SEOAgency,
    myReport: SEOReport,
    localReports: list,
    searchReports: list,
    siteAuditReports: list,
    clientName: str = "SEO Report",
) -> str | None:
    tpl = DocxTemplate("SEO/Resources/templates/master_template.docx")
    basictpl = DocxTemplate("SEO/Resources/templates/master_template.docx")
    reportCover: Subdoc = tpl.new_subdoc(myReport.fileName)
    basicCover: Subdoc = basictpl.new_subdoc(myReport.fileName)
    has_subs = True
    subs: list = []
    basicsubs: list = []
    outputPath: str = f"{agency.output_directory}/reports/{cleanFilename(clientName)}"
    if not os.path.exists(outputPath):
        os.makedirs(outputPath)
    basicPath: str = f"{outputPath}/basic_report_{time.strftime('%b_%d_%y_%H%M%S', time.gmtime())}.docx"
    outputPath += (
        f"/master_report_{time.strftime('%b_%d_%y_%H%M%S', time.gmtime())}.docx"
    )
    report: str
    for report in localReports:
        if report is None or not os.path.exists(report):
            continue
        subs.append(tpl.new_subdoc(report))
        basic = report.replace(".docx", "basic.docx")
        if os.path.exists(basic):
            basicsubs.append(basictpl.new_subdoc(basic))
    for report in searchReports:
        if report is None or not os.path.exists(report):
            continue
        subs.append(tpl.new_subdoc(report))
        basic = report.replace(".docx", "basic.docx")
        if os.path.exists(basic):
            basicsubs.append(basictpl.new_subdoc(basic))
    for report in siteAuditReports:
        if report is None or not os.path.exists(report):
            continue
        subs.append(tpl.new_subdoc(report))
    context = {
        "agencyName": agency.company_name,
        "agencyPhone": agency.company_phone,
        "clientName": clientName,
        "footerLeft": myReport.footerLeft,
        "footerCenter": myReport.footerCenter,
        "footerRight": myReport.footerRight,
        "reportCoverPage": reportCover,
        "has_subs": has_subs,
        "subs": subs,
    }
    basiccontext = {
        "agencyName": agency.company_name,
        "agencyPhone": agency.company_phone,
        "clientName": clientName,
        "footerLeft": myReport.footerLeft,
        "footerCenter": myReport.footerCenter,
        "footerRight": myReport.footerRight,
        "reportCoverPage": basicCover,
        "has_subs": has_subs,
        "subs": basicsubs,
    }
    try:
        tpl.render(context)
        tpl.save(outputPath)
        basictpl.render(basiccontext)
        basictpl.save(basicPath)
        return outputPath
    except:
        return None


def createSitemapAuditReport(auddict: dict) -> str | None:
    tpl = DocxTemplate("SEO/Resources/templates/website_sitemap_audit_template.docx")
    context = dict()
    context["baseURL"] = auddict["baseURL"]
    if not "clientName" in auddict:
        context["clientName"] = "False"
    else:
        context["clientName"] = auddict["clientName"]
    if not "clientCompanyName" in auddict:
        context["clientCompanyName"] = "False"
    else:
        context["clientCompanyName"] = auddict["clientCompanyName"]
    context["clientName"].replace("&", "and").replace(">", "").replace("<", "")
    context["clientCompanyName"].replace("&", "and").replace(">", "").replace("<", "")
    if os.path.exists(auddict["generalImage"]):
        context["generalImage"] = InlineImage(
            tpl, auddict["generalImage"], width=Inches(3)
        )
    else:
        context["generalImage"] = InlineImage(tpl, "blank.png", width=Inches(3))
    if os.path.exists(auddict["errorsimg"]):
        context["errorsimg"] = InlineImage(tpl, auddict["errorsimg"], width=Inches(3))
    else:
        context["errorsimg"] = InlineImage(tpl, "blank.png", width=Inches(3))
    context["criticalerrors"] = str(auddict["criticalerrors"])
    context["moderate"] = str(auddict["moderateerrors"])
    context["minor"] = str(auddict["minorerrors"])
    context["suggestions"] = str(auddict["suggestions"])
    context["tableimg"] = InlineImage(
        tpl, auddict["tableimg"], width=Inches(6.5), height=Inches(8)
    )
    context["auditReports"] = list()
    for report in auddict["auditReports"]:
        tdic = dict()
        tdic["title"] = report["title"]
        tdic["url"] = report["url"]
        if os.path.exists(report["image"]):
            tdic["image"] = InlineImage(tpl, report["image"], width=Inches(3))
        else:
            tdic["image"] = InlineImage(tpl, "blank.png", width=Inches(3))
        tdic["errorTextcritical"] = str(report["errorTextcritical"])
        tdic["errorTextmoderate"] = str(report["errorTextmoderate"])
        tdic["errorTextminor"] = str(report["errorTextminor"])
        tdic["errorTextdonthurt"] = str(report["errorTextdonthurt"])
        tdic["hascomlist"] = report["hascomlist"]
        tdic["comlist"] = list()
        if tdic["hascomlist"]:
            for key, val in report["comlist"].items():
                tdic["comlist"].append({"key": key, "val": val})
        else:
            tdic["comlist"].append({"key": "", "val": ""})
        context["auditReports"].append(tdic)
    try:
        tpl.render(context)
        outpath: str = f"{BackSEOSettings().agencyOutput}/reports/{cleanFilename(context['clientName'])}/sitemaps"
        if not os.path.exists(outpath):
            os.makedirs(outpath)
        savepath: str = str(
            f"{outpath}/audit_report{time.strftime('%b_%d_%y', time.gmtime())}.docx"
        )
        tpl.save(savepath)
        return savepath
    except Exception as e:
        # print(e)
        return None


def createSearchReport(report) -> str | None:
    tpl = DocxTemplate("SEO/Resources/templates/search_audit_template.docx")
    basictpl = DocxTemplate("SEO/Resources/templates/search_audit_template_basic.docx")
    context = dict()
    basiccontext = dict()
    if "&" in report["client"]:
        report["client"] = str(report["client"]).replace("&", "and")
    if ">" in report["client"]:
        report["client"] = str(report["client"]).replace(">", "")
    if "<" in report["client"]:
        report["client"] = str(report["client"]).replace("<", "")
    context["keyword"] = report["keyword"]
    context["client"] = report["client"]
    basiccontext["keyword"] = report["keyword"]
    basiccontext["client"] = report["client"]
    context["tblimg"] = InlineImage(
        tpl, report["tblimg"], width=Inches(6), height=Inches(7.5)
    )
    basiccontext["tblimg"] = InlineImage(
        basictpl, report["tblimg"], width=Inches(6), height=Inches(7.5)
    )
    context["kwstats"] = report["kwstats"]
    # context["kwimg"] = InlineImage(tpl, report["kwimg"], width=Inches(6), height=Inches(8))
    try:
        tpl.render(context)
        outpath: str = f"{BackSEOSettings().agencyOutput}/reports/{cleanFilename(report['client'])}/searches/{report['keyword']}"
        if not os.path.exists(outpath):
            os.makedirs(outpath)
        savepath: str = (
            f"{outpath}/search_report{time.strftime('%b_%d_%y', time.gmtime())}.docx"
        )
        basic: str = f"{outpath}/search_report{time.strftime('%b_%d_%y', time.gmtime())}basic.docx"
        tpl.save(savepath)
        basictpl.render(basiccontext)
        basictpl.save(basic)
        return savepath
    except Exception as e:
        print(e)
        return None


if __name__ == "__main__":
    myAgency = SEOAgency()
    myReport = SEOReport(myAgency)
    myReport.coverPage()
    myReport.coverHeaderFooter()
    myReport.save("SomeFile")
    tpl = DocxTemplate("SEO/Resources/templates/master_template.docx")
    reportCover = tpl.new_subdoc(myReport.fileName)
    has_subs = True
    subs = []
    newR = SEOReport(
        SEOAgency(
            "Sub Agency",
            "321 Sub St. City, State 54321",
            "123-456-7890",
            "subemail@submail.sub",
            "subwebsite.sub",
            "bSEOLogo.png",
        )
    )
    newR.coverPage()
    newR.save("SubFile")
    subs.append(tpl.new_subdoc(newR.fileName))
    context = {
        "reportCoverPage": reportCover,
        "has_subs": has_subs,
        "subs": subs,
    }
    tpl.render(context)
    tpl.save("output/reports/SEO Report with Subs.docx")
